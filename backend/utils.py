# utils.py

"""Utility functions for web scraping and cover letter text processing."""

import copy
import requests

from bs4 import BeautifulSoup
from docx import Document
from docx.text.paragraph import Paragraph

from config import NAMESPACE, PLACEHOLDER_PATTERN


def scrape_url(url: str):
    """Fetch and extract text content from a web page by URL."""
    try:
        headers = {"User-Agent": "CoverLetterApp/1.0"}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        return text[:5000] if len(text) > 100 else None
    except Exception:
        return None


def clean_cover_letter(text: str) -> str:
    """Remove service text before and after the actual letter."""
    start_markers = ["Dear"]
    for marker in start_markers:
        idx = text.find(marker)
        if idx != -1:
            text = text[idx:]
            break

    end_markers = ["Kind regards, Alina", "Kind regards,\nAlina"]
    for marker in end_markers:
        idx = text.find(marker)
        if idx != -1:
            text = text[:idx + len(marker)]
            break

    return text.strip()


def find_placeholders(file_path: str) -> list:
    """Find all {{PLACEHOLDER}} patterns in a DOCX file."""
    doc = Document(file_path)
    plaсeholders = set()

    for paragraph in _iter_all_paragraphs(doc):
        plaсeholders.update(PLACEHOLDER_PATTERN.findall(paragraph.text))

    return sorted(plaсeholders)


def fill_template(file_path: str, data: dict, output_path: str,
                  multiline: bool = True):
    """Replace {{PLACEHOLDER}} patterns in a DOCX with actual content."""
    doc = Document(file_path)
    paragraphs = list(_iter_all_paragraphs(doc))

    for paragraph in paragraphs:
        if "{{" not in paragraph.text:
            continue
        for key, value in data.items():
            placeholder = "{{" + key + "}}"
            if placeholder not in paragraph.text:
                continue


    doc.save(output_path)


def _iter_all_paragraphs(doc):
    """Yield all paragraphs from the document and its tables."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _replace_with_text(paragraph: Paragraph, placeholder: str, replacement: str) -> None:
    """Replace a placeholder with plain text, preserving original formatting."""
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            return

    if placeholder in paragraph.text and paragraph.runs:
        paragraph.runs[0].text = paragraph.text.replace(placeholder, replacement)
        for run in paragraph.runs[1:]:
            run.text = ""
